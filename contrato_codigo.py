import tkinter as tk
from tkinter import messagebox
from datetime import datetime
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words

def obtener_nombre_mes(numero_mes):
    meses = {
        1: "enero",
        2: "febrero",
        3: "marzo",
        4: "abril",
        5: "mayo",
        6: "junio",
        7: "julio",
        8: "agosto",
        9: "septiembre",
        10: "octubre",
        11: "noviembre",
        12: "diciembre"
    }
    return meses.get(numero_mes, "error con meses nombre")
def suma_meses(fecha_ingresada,meses_num):
    fecha_objeto = datetime.strptime(fecha_ingresada, "%d/%m/%Y")
    fecha_nueva = fecha_objeto + relativedelta(months=meses_num)
    return fecha_nueva.strftime("%m/%Y")


root = tk.Tk() #inicia la interfaz
root.title("Pytrato :D")

root.geometry('600x400') #Tamaño de interfaz


label = tk.Label(root, text="NOMBRE")
label.pack()

NOMBRE = tk.Entry(root)
NOMBRE.pack()

label = tk.Label(root, text="DOCUMENTO")
label.pack()

DNI = tk.Entry(root)
DNI.pack()

label = tk.Label(root, text="PISO (PRIMER, SEGUNDO, TERCERO, etc)")
label.pack()

PISO = tk.Entry(root)
PISO.pack()

label = tk.Label(root, text="CUARTO (1, 2, 3, etc)")
label.pack()

CUARTO = tk.Entry(root)
CUARTO.pack()

label = tk.Label(root, text="ALQUILER (450, 520, 550, etc)")
label.pack()

ALQ_SOLES = tk.Entry(root)
ALQ_SOLES.pack()

label = tk.Label(root, text="FECHA DE INICIO (00/00/0000)")
label.pack()

fecha_inicio = tk.Entry(root)
fecha_inicio.pack()

label = tk.Label(root, text="GARANTIA (450, 520, 550, etc)")
label.pack()

GAR_SOL = tk.Entry(root)
GAR_SOL.pack()

label = tk.Label(root, text="DURACION/MESES (1, 2, 5, 12)")
label.pack()

DURACION = tk.Entry(root)
DURACION.pack()

#Codigo ejecutable al presionar el boton de ingreso
def pytrato():
    try:
        obtener_dato()
    except:
        messagebox.showerror("Error", "Vuelva a abrir el programa e intente denuevo")
        root.destroy()

def obtener_dato():
    global NOMBRE
    global DNI
    global PISO
    global CUARTO
    global ALQ_SOLES
    global fecha_inicio
    global GAR_SOL
    global DURACION
    DURACION = DURACION.get()
    NOMBRE = NOMBRE.get()
    DNI = DNI.get()
    PISO = PISO.get()
    CUARTO = CUARTO.get()
    ALQ_SOLES= ALQ_SOLES.get()
    ALQUILER = (num2words(int(ALQ_SOLES), lang='es')).upper()
    fecha_inicio = fecha_inicio.get()
    dia, mes1, año = fecha_inicio.split("/")
    mes = (obtener_nombre_mes(int(mes1)))
    nuevo_tiempo = suma_meses(fecha_inicio, int(DURACION)) #genera la fecha final
    mes_final1, año_final = nuevo_tiempo.split("/")
    mes_final = (obtener_nombre_mes(int(mes_final1)))

    GAR_SOL = GAR_SOL.get()
    GARANTIA = (num2words(int(GAR_SOL), lang='es')).upper()

    ###############################################################
    def read_docx(docx_file):
        """
        Lee el documento y lo convierte en una lista de parrafos para su modificacion
        """
        doc = Document(docx_file)
        text = ""
        for i in doc.paragraphs:
            text += i.text + "\n"
        return text

    # Uso de la función read_docx
    documento = 'testing_save.docx'
    contenido = read_docx(documento)

    # print("Contenido del documento:")
    # print(contenido)
    var1 = contenido.split("\n")
    # print(var1)


    doc = Document()
    for parrafo in var1:
        if "[NOMBRE]" or "[DNI]" or "[PISO]" or "[CUARTO]" or "[ALQUILER]" or "[ALQ_SOLES]" or "[dia]" or "[mes]" or "[año]" or "[mes_final]" or "[año_final]" or "[GARANTIA]" or "[GAR_SOL]" in parrafo:
            parrafo = parrafo.replace("[NOMBRE]", NOMBRE.upper()).replace("[DNI]", DNI).replace("[PISO]", PISO).replace(
                "[CUARTO]", CUARTO).replace("[ALQUILER]", ALQUILER).replace("[ALQ_SOLES]", ALQ_SOLES).replace("[dia]",
                                                                                                              dia).replace(
                "[mes]", mes).replace("[año]", año).replace("[mes_final]", mes_final).replace("[año_final]",
                                                                                              año_final).replace(
                "[GARANTIA]", GARANTIA).replace("[GAR_SOL]", GAR_SOL).replace("[DURACION]", DURACION)
            contrato = ""
            contrato += parrafo + ""
            if contrato == "":
                pass
            else:
                doc.add_paragraph(contrato)

    doc.save(f"{NOMBRE}Pytrato.docx")

    ###################
    def cambiar_justificado(docx_file):
        """
        Cambia el justificado
        """
        doc = Document(docx_file)

        for paragraph in doc.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.save(docx_file)

    documento = f"{NOMBRE}Pytrato.docx"

    cambiar_justificado(documento)
    ##############################################################





    root.destroy()
    print("El dato ingresado es:", NOMBRE)


# Crear un botón
boton = tk.Button(root, text="Obtener dato", command=pytrato)
boton.pack()

# Iniciar el bucle de eventos de Tkinter
root.mainloop()